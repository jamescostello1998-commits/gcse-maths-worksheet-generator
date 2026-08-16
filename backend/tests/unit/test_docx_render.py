"""Tests for the Word (.docx) renderers (app/docx/render.py) and the format
branch on the worksheet / modelled-example routes.

Structural assertions on the produced .docx (title, per-question paragraphs,
the Worked Solutions / Answers sections, embedded diagram images, and real
native <m:oMath> equation elements) - the docx equivalent of the visual QA
done by rendering in real Word, which a unit test can't do.
"""

import io
import re
import random

from docx import Document
from fastapi.testclient import TestClient

from app.core.registry import get_topic
from app.docx.render import render_modelled_example_docx, render_worksheet_docx
from app.main import app
from app.worksheet.builder import build_worksheet

_DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# A fractional-power topic (native superscripts + stacked fractions), a plain
# (non-hoisted) fraction topic, and a diagram-bearing topic.
# NB: several other fractions.py topics are hoisted (see app/worksheet/
# builder.py's HOISTED_VERB_INSTRUCTIONS) - deliberately not used here, since
# this file's "plain worksheet" structural assertions assume an un-hoisted
# Q1./Q2./... layout. fractions_simplify_F is not in that registry.
_FRAC_POWER_TOPIC = "algebraic_indices_H"
_FRACTION_TOPIC = "fractions_simplify_F"
_DIAGRAM_TOPIC = "area_triangle_F"
_VECTOR_TOPIC = "vectors_arithmetic_H"
_GEOMETRIC_VECTOR_TOPIC = "geometric_vectors_H"


def _worksheet_doc(topic_id: str, *, count: int = 6, answers_only: bool = False) -> Document:
    topic = get_topic(topic_id)
    ws = build_worksheet(topic_id, topic.fixed_tier, count=count, rng=random.Random(1))
    data = render_worksheet_docx(ws, answers_only=answers_only)
    assert data[:2] == b"PK"  # a .docx is a zip archive
    return Document(io.BytesIO(data))


def _paragraph_texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs]


def _question_paragraph_count(doc: Document) -> int:
    return sum(1 for p in doc.paragraphs if re.match(r"^Q\d+\.", p.text))


def test_worksheet_title_is_the_topic_name():
    topic = get_topic(_FRACTION_TOPIC)
    doc = _worksheet_doc(_FRACTION_TOPIC)
    assert doc.paragraphs[0].text == topic.display_name


def test_worksheet_has_one_paragraph_per_question_and_worked_solutions():
    doc = _worksheet_doc(_FRACTION_TOPIC, count=6)
    # Six on the question page + six under Worked Solutions (each headed "Q{n}").
    assert _question_paragraph_count(doc) == 6
    assert "Worked Solutions" in _paragraph_texts(doc)
    assert "Answers" not in _paragraph_texts(doc)


def test_answers_only_renders_an_answers_list_not_worked_solutions():
    doc = _worksheet_doc(_FRACTION_TOPIC, count=6, answers_only=True)
    texts = _paragraph_texts(doc)
    assert "Answers" in texts
    assert "Worked Solutions" not in texts


def test_fraction_topic_produces_native_fraction_equations():
    doc = _worksheet_doc(_FRACTION_TOPIC)
    xml = doc.element.xml
    assert "oMath" in xml
    assert "}f>" in xml or ":f>" in xml  # an <m:f> stacked-fraction element


def test_fractional_power_topic_produces_native_superscript_equations():
    doc = _worksheet_doc(_FRAC_POWER_TOPIC)
    xml = doc.element.xml
    assert "oMath" in xml
    assert "sSup" in xml  # a native superscript (x^2, (25x^4)^(1/2), etc.)


def test_diagram_topic_embeds_at_least_one_image():
    doc = _worksheet_doc(_DIAGRAM_TOPIC)
    assert len(doc.inline_shapes) >= 1


def test_vector_topic_produces_native_column_vector_equations():
    doc = _worksheet_doc(_VECTOR_TOPIC)
    xml = doc.element.xml
    assert "oMath" in xml
    # A native column vector is an <m:d> delimiter around an <m:m> matrix.
    assert "}d>" in xml or ":d>" in xml
    assert "}m>" in xml or ":m>" in xml


def test_vector_letters_render_bold():
    # \vec{a}/\vec{b} must be bold (matching the PDF's <b> treatment), while
    # ordinary prose "a" (e.g. "is a triangle") stays non-bold.
    doc = _worksheet_doc(_GEOMETRIC_VECTOR_TOPIC)
    bold_vector_runs = [
        run
        for para in doc.paragraphs
        for run in para.runs
        if run.text in {"a", "b"} and run.bold and run.font.name == "Cambria Math"
    ]
    assert bold_vector_runs, "expected at least one bold Cambria Math vector letter"


def test_modelled_example_docx_structure():
    topic = get_topic(_FRACTION_TOPIC)
    rng = random.Random(2)
    example = topic.generate_modelled_example(topic.fixed_tier, rng)
    practice = build_worksheet(_FRACTION_TOPIC, topic.fixed_tier, count=5, rng=rng)
    data = render_modelled_example_docx(
        topic.display_name, topic.fixed_tier, example, practice.questions, topic.preamble_lines or ()
    )
    doc = Document(io.BytesIO(data))
    texts = _paragraph_texts(doc)
    assert doc.paragraphs[0].text == topic.display_name
    assert "How it works" in texts
    assert "Now You Try" in texts
    # Five faded practice questions.
    assert _question_paragraph_count(doc) == 5


# --- route format branch ---------------------------------------------------


def test_route_worksheet_docx_returns_word_document():
    client = TestClient(app)
    r = client.post(
        "/api/worksheets",
        json={"topic_id": _FRAC_POWER_TOPIC, "tier": "higher", "count": 6, "format": "docx"},
    )
    assert r.status_code == 200
    assert r.headers["content-type"] == _DOCX_MEDIA_TYPE
    assert r.headers["content-disposition"].endswith('.docx"')
    assert r.content[:2] == b"PK"


def test_route_modelled_example_docx_returns_word_document():
    client = TestClient(app)
    r = client.post(
        "/api/modelled-examples",
        json={"topic_id": _FRAC_POWER_TOPIC, "tier": "higher", "format": "docx"},
    )
    assert r.status_code == 200
    assert r.headers["content-type"] == _DOCX_MEDIA_TYPE
    assert r.headers["content-disposition"].endswith('.docx"')


def test_route_default_format_is_still_pdf():
    client = TestClient(app)
    r = client.post("/api/worksheets", json={"topic_id": _FRAC_POWER_TOPIC, "tier": "higher"})
    assert r.status_code == 200
    assert r.headers["content-type"] == "application/pdf"
    assert r.headers["content-disposition"].endswith('.pdf"')
